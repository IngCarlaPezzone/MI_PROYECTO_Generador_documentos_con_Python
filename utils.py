import pandas as pd
from docxtpl import DocxTemplate
import docx
from copy import deepcopy

### FUNCIONES DE UTILIDAD

def delete_paragraph(paragraph):
    '''
    Esta función es para borrar un determinado parrafo en el documento.
    '''
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = p._element = None

def insert_paragraph_after(paragraphs, idx, text=None):
    '''
    Función que inserta un parrafo a continuación de un determinado lugar en el documento.
    '''
    next_paragraph_idx = idx + 1
    if idx == len(paragraphs):
        return documento.add_paragraph(text)
    next_paragraph = paragraphs[next_paragraph_idx]
    return next_paragraph.insert_paragraph_before(text)

def reemplazar_parrafo(document, palabra, nuevo_texto):
    '''
    Esta función sirve para buscar una palabra determinada en el documento
    y reemplazarla por un nuevo_texto.
    Actualmente no se usa, pero la deje por si es de utilidad.
    '''
    for paragraph in document.paragraphs:
            if palabra in paragraph.text:
                print (paragraph.text)
                paragraph.text = nuevo_texto


def excel_to_word(titulo, archivo, codigo):
    '''
    Define de dónde viene el nuevo texto que se va a actualizar en el documento.
    Lo busca en el excel por el código del documento y selecciona el texto según el titulo del documento.
    '''
    # ACCEDEMOS A LA INFORMACION DE CADA FILA 
    fila = (archivo[archivo.codigo == codigo])   
    # ACCEDEMOS A LA INFORMACION DE CADA CELDA
    texto = fila [titulo].tolist()
    texto = str(texto[0])
    return texto

def generador_doc (doc, archivo):
    ''' 
    Esta función lo que hace es primero genera un diccionario por cada fila
    del excel donde esta toda la información y la guarda en dic_record.
    Luego toma los datos del documento que no estaba creado previamente y lo guarda en documento.
    Finalmente crea el documento a partir de la plantilla.
    '''
    # GENERA UN DICCIONARIO VACIO QUE CONTENDRA POR CADA CODIGO 
    # EL DICCIONARIO DE CADA FILA DE INFORMACIÓN
    dic_record = {}
    # CREA UN DICCIONARIO CODIGO-INFORMACIÓN
    # DONDE INFORMACIÓN ES LO QUE ESTA EN EXCEL POR CADA FILA
    for record in archivo.to_dict(orient="records"):
        codigo = record.get('codigo')
        dic_record[codigo] = record
        
    #DOCUMENTO ES EL DICCIONARIO INDIVIDUAL DE INFORMACIÓN DE UNA SOLA FILA
    documento = dic_record.get(doc)
    
    codigo = documento['codigo']
    nombre = documento['nombre']
    
    # TOMA LA PLANTILLA Y REPARTE LA INFORMACIÓN
    plantilla = DocxTemplate('./Plantilla.docx')
    plantilla.render(documento)

    # GUARDA EL DOCUMENTO CON ESTE FORMATO DE TITULO
    plantilla.save('./documentos/{}-{}.docx'.format(codigo, nombre))

def actualizacion(nombre_doc, archivo):
    '''
    Es la función que permite actualizar la información de los documentos existentes, 
    cambiando parrafos específicos, tomando la información del excel.
    '''
    codigo = nombre_doc[:5] # SI EL CÓDIGO CAMBIA DE FORMATO LL-NN, HAY QUE REVISAR ACA
    # DEFINO EL PATH DEL DOCUMENTO A ACTUALIZAR PARA QUE LO PISE
    path_doc = './documentos/' + nombre_doc
    # GENERO EL OBJETO DEL DOCUMENTO
    documento = docx.Document(path_doc)
    # GENERO UNA LISTA CON TODOS LOS OBJETOS DEL WORD
    paragraphs = list(documento.paragraphs)
    # HAGO UNA LISTA VACÍA PARA RECORRER AGREGAR LOS TEXTOS DE LOS OBJETOS
    fullText = []
    # RECORRO LA LISTA PARA OBTENER LOS INDICES DE LOS ENCABEZADOS
    for para in paragraphs:
        parrafo = para.text
        fullText.append(parrafo)

    #DEFINO LOS INDICES DE UBICACIÓN DE LOS ENCABEZADOS
    encabezado_objetivos = fullText.index('OBJETIVOS')
    encabezado_responsabilidades = fullText.index('RESPONSABILIDADES')
    encabezado_RegAso = fullText.index('DOCUMENTOS ASOCIADOS')
    encabezado_descrip = fullText.index('DESCRIPCIÓN')

    # DEFINO LOS RANGOS DE TEXTO INTERMEDIOS ENTRE LOS ENCABEZADOS
    descripcion_objetivos = documento.paragraphs[encabezado_objetivos+1:encabezado_responsabilidades]
    descripcion_responsabilidades = documento.paragraphs[encabezado_responsabilidades+1:encabezado_descrip]
    descripcion_RegAso = documento.paragraphs[encabezado_RegAso+1:]

    # GENERO UNA LISTA CON LOS RANGOS DE TEXTOS EXISTENTES
    lista_textos_existentes = [descripcion_objetivos, descripcion_responsabilidades, descripcion_RegAso]

    # RECORRO EL RANGO DE TEXTO EXISTENTE EN OBJETIVOS PARA BORRAR EL CONTENIDO
    for textos_existentes in lista_textos_existentes:
        for parrafos in textos_existentes:
            borrar_textos = delete_paragraph(paragraph=parrafos)

    # INSERTO UN DETERMINADO TEXTO DEBAJO DE CADA ENCABEZADO
    paragraphs = list(documento.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text == 'OBJETIVOS':
            insert_paragraph_after(paragraphs, idx, excel_to_word('objetivos', archivo, codigo))
        if paragraph.text == 'RESPONSABILIDADES':
            insert_paragraph_after(paragraphs, idx, excel_to_word('responsabilidades', archivo, codigo))
            
    documento.save(path_doc)

def add_hyperlink(paragraph, url, text, color, underline):
    """
    Fuente: https://github.com/python-openxml/python-docx/issues/384
    Función que coloca un hipervínculo dentro de un objeto párrafo.

    :param párrafo: El párrafo al que vamos a añadir el hipervínculo.
    :param url: Una cadena que contiene la url requerida.
    :param texto: El texto que se muestra para la url
    :return: El objeto hipervínculo
    """
    # ESTO OBTIENE EL ACCESO AL ARCHIVO document.xml.rels Y
    # OBTIENE UN NUEVO VALOR DE IDENTIFICACIÓN DE LA RELACIÓN
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # CREA LA ETIQUETA w:hyperlink Y AÑADE LOS VALORES NECESARIOS
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # CREA UN ELEMENTO w:r 
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # CREA UN NUEVO ELEMENTO w:rPr
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # AGREGA COLOR SI ES QUE SE DA EL MISMO
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # ELIMITA EL SUBRAYADO SI SE SOLICITA
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # UNE TODOS LOS ELEMENTOS xml, AÑADE EL TEXTO REQUERIDO AL ELEMENTO w:r
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def agrega_hipervinculos(document, archivo, doc_exist, lista_list, lista2):
    '''
    Esta función prepara el texto que se tiene que agregar al final del documento y luego
    le genera el hiperlink con la función add_hyperlink.
    Se hizo de esta forma, porque al momento de cerrar este scrip no había forma de que el
    hiperlink se coloque en un lugar específico. Solo se puede poner al final del documento.
    '''
    # CON ESTE BUCLE LO QUE SE HACE ES RECORRER LOS CÓDIGOS QUE ESTAN EN LA lista_list
    # PARA QUE SI ESE CÓDIGO COINCIDE CON LOS PRIMEROS 5 CARACTERIRES DEL DOCUMENTOS EXISTENTE
    # ENTONCES BUSQUE LA INFORMACIÓN QUE LE CORRESPONDE A DOCUMENTOS ASOCIADOS PARA ESE CÓDIGO
    # Y LOS VA AGREGANDO A UNA LISTA (lista_p).

    for codigo in lista_list:
        lista_p = []
        if codigo == doc_exist[:5]: # SI EL CÓDIGO CAMBIA DE FORMATO LL-NN, HAY QUE REVISAR ACA
            parrafos = excel_to_word('doc_asociados', archivo, codigo)
            lista_p.append(parrafos) #lista_p es una lista del tipo ['RE-01\nRE-02\nRE-03']
            # CON EL TRY LO QUE SE BUSCA ES EVITAR Y AVISAR ERRORES CUANDO UN DOCUMENTO ASOCIADO
            # FUE NOMBRADO EN EL EXCEL PERO NO SE GENERÓ SU CORRESPONDIENTE FILA.
            # SI PASA ESO, EL except YA A DECIR CUAL ES EL DOCUMENTO NO TIENE SU FILA
            # SI NO ESTA, NO TERMINA EL PROCESO DE AGREGAR LOS LINK Y ALGUNOS DOCUMENTOS PUEDEN NO
            # TENER SU LINK
            try:
                # CON ESTE BUCLE RECORRO LA LISTA lista_p Y SI ES DISTINTO A "No tiene documentos asociados."
                # ENTONCES ME SPLITEA EL TEXTO DE lista_p CUANDO TIENE \n
                # UBICA ESE VALOR EN LA lista_list Y DEVUELVE SI index PARA ENCONTRAR SU DESCRIPCIÓN
                # EN LA lista2
                # FINALMENTE FORMA item QUE ES EL CODIGO + DESCRIPCIÓN QUE ES EL TEXTO QUE QUEREMOS 
                # QUE SE AGREGUE AL FINAL DEL DOCUMENTO
                for doc in lista_p:
                    if doc != "No tiene documentos asociados.":
                        for d in doc.split():
                            ubicacion = lista_list.index(d)
                            descrip = lista2[ubicacion]
                            # GENERA EL CÓDIGO CON EL DOCUMENTO ASOCIADO
                            item = d + '-' + descrip
                            # SE SETEA EL TEXTO A AGREGAR, LA UBICACIÓN DEL DOCUMENTO A LINKIAR, EL TEXTO QUE MUESTRA EL LINK Y EL COLOR
                            add_hyperlink(document.add_paragraph(item), item + '.docx', ' (Ir al documento)', '0000FF', False)
                            # SE GUARDA PISANDO EL MISMO DOCUMENTO
                            document.save('./Documentos/' + doc_exist)

                    else:
                        # SI NO TIENE DOCUMENTOS ASOCIADOS SOLO COLOCA UNA LEYENDA Y NO PONE HIPERLINK
                        document.add_paragraph("No tiene documentos asociados.")
                        document.save('./Documentos/' + doc_exist)
            except:
                print('No se encontró el Documento Asociado {}. Agregarlo al Excel para completar el proceso'.format(d))